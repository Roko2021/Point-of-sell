import tkinter
from tkinter import *
from tkinter import messagebox
import datetime
import os
from collections import defaultdict
from docx import Document
import csv
import win32print
import win32api
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
def seller1():
    def multiply_price(name):
        with open('products.csv', 'r', newline='', encoding='utf-8') as file:
            # row_number = 1
            reader = csv.DictReader(file)
            for row in reader:
                # product_name = row['prodects']  # استخدم اسم العمود الصحيح هنا
                price = row['price']
        # try:
        
        value = float(entries[name].get())# القيمة المدخلة في مربع الإدخال
        price = float(row['price'])  # السعر
        result = value * price  # الضرب
        fatora_listbox.insert(END,f'             {name}                                      {value}                                {result}\n')
        # se3r = 
        total_se3r1 = float(total_fatora.get("1.0", "end-1c"))
        total_se3r = total_se3r1 + result
        total_fatora.delete('1.0',END)
        total_fatora.insert(END,f'{total_se3r}')

    def delete_food1():
        
        selected_indices = fatora_listbox.curselection()
        selected_items = [fatora_listbox.get(index) for index in selected_indices]
        
        # قطع النص أو الأرقام من العناصر المحددة
        if selected_items:
            fatora_listbox.delete(selected_indices)
            # هنا يمكنك استخدام عملية القطع (slicing) لاستخراج جزء معين من النص أو الأرقام
            part_of_text = selected_items[0][100:120]  # على سبيل المثال، نأخذ أول عنصر ونقوم بقص أول 5 حروف منه
            part_of_text1 = float(part_of_text)
            total_se3r1 = float(total_fatora.get("1.0", "end-1c"))
            total_se3r = total_se3r1 - float(part_of_text1)
            total_fatora.delete('1.0',END)
            total_fatora.insert(END,f'{total_se3r}')

    def close():
        root.destroy()
        
    def fatora_empty():
        fatora_listbox.delete(0,END)
        total_fatora.delete('1.0',END)
        total_fatora.insert(END,0)
        total_fatora.tag_configure("center", justify='center')
        total_fatora.tag_add("center", "1.0", "end")
        
    def save_fatora():
        
        today_date = datetime.datetime.now().strftime("%Y-%m-%d")
        today_date1 = datetime.datetime.now().strftime("%Y-%m")
        
        folder_path = os.path.join(os.getcwd(),'daily',today_date1, today_date)
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
            
        file_name = os.path.join(folder_path, f"{len(os.listdir(folder_path)) + 1}.csv")
        file_name1 = os.path.join(folder_path, "ta2rer.csv")
        file_name2 = os.path.join(folder_path, "save.csv")
        
        data_save = []
        if fatora_listbox.size() != 0:
            with open(file_name, "w",encoding='utf-8') as file:
                write1 = csv.DictWriter(file,fieldnames=['المنتجات','العدد','السعر','الاجمالى'])
                write1.writeheader()
            
            for row in range(fatora_listbox.size()):
                mount = fatora_listbox.get(row).split()
                data_save.append(mount)
            

        else:
            messagebox.showinfo('error','لا يوجد فواتير لتسجيلها')
            
        total_test = total_fatora.get('1.0',END)
        with open(file_name, "w",encoding='utf-8') as file5:
            writer = csv.DictWriter(file5,fieldnames=['المنتجات','العدد','السعر','الاجمالى'])
            writer.writerow({'المنتجات':'المنتجات','العدد':'العدد','السعر':'السعر','الاجمالى':'الاجمالى'})
            for row in data_save:
                writer.writerow({'المنتجات':row[0]+ ' ' +row[1],'العدد':row[2],'السعر':row[3]})
                
            writer.writerow({'الاجمالى':total_test})
        file_txt = os.path.join(folder_path, f"{len(os.listdir(folder_path))}.txt")
        file_docx = os.path.join(folder_path, f"{len(os.listdir(folder_path))}.docx")
        total_madfo3 = total_fatora1.get()
        total_ba2y = total_fatora2.get('1.0',END)
        #نقل البيانات من ملف CSV الى ملف txt
        #فتح ملف CSV وتسجيل بياناته 
        #فتح الملف النصى للكتابه وطباعته
        # with open(file_txt,'w',encoding='utf-8') as file11:
        #         file11.write('العدد\tالمنتج\t\tالسعر\n')
        #         for row in data_save:
        #             # row = '\t'.join(row)
        #             file11.write(row[2]+ '\t'+row[0] + ' ' + row[1] + '\t' + row[3] + '\n')
        #         file11.write( '\n'+'الاجمالى\t' + total_test )
        #         file11.write('المدفوع\t' + total_madfo3 + '\n')
        #         file11.write('الباقى\t' + total_ba2y )
                
        doc = Document()

        header = doc.add_heading('الليثى', 0)

        p = doc.add_paragraph('الليثى حواوشى فاتوره')
        p1 = doc.add_paragraph(':التاريخ')

        table_header = ['العدد', 'المنتج', 'السعر']
        table = doc.add_table(rows=1, cols=3)

        # إضافة عناوين العمود مرة واحدة في الصف الأول
        hdr_cells = table.rows[0].cells
        for i in range(3):
            hdr_cells[i].text = table_header[i]
            hdr_cells[i].width = 1

        hanco = []
        for row in data_save:
            hanco.append([row[2],[row[0] , ' ', row[1]], [row[3]]])

        for row_data in hanco:
            row_cells = table.add_row().cells
            for i in range(3):
                row_cells[i].text = row_data[i]
                row_cells[i].width = 1
        
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
        doc.save(file_docx)
        #امر الطباعه
        win32api.ShellExecute(0, 'print', file_docx, None,'.',1)
        ta2rer_today = os.path.join(folder_path, "ta2rer_today.csv")
        if not os.path.exists(ta2rer_today):
            montag1 = []
            with open('products.csv', 'r', newline='', encoding='utf-8') as file:
                reader = csv.DictReader(file)
                for row in reader:
                    montag1.append(row['prodects'])
            # print(montag1)
            with open(ta2rer_today,'w',encoding='utf-8') as file9:
                write4 = csv.DictWriter(file9,fieldnames=['المنتجات','العدد الكلى','السعر','اجمالى اليوم'])
                write4.writeheader()
                for row in montag1:
                    write4.writerow({'المنتجات':row,'العدد الكلى':0,'السعر':0})
                write4.writerow({'اجمالى اليوم':0})
        gam3_3dd1 = pd.read_csv(file_name)
        gam3_3dd2 = pd.read_csv(ta2rer_today)
        merge_df = pd.merge(gam3_3dd1, gam3_3dd2, on='المنتجات',how='right')
        merge_df['العدد الكلى'] = merge_df['العدد'] + merge_df['العدد الكلى']
        merge_df['اجمالى اليوم'] = merge_df['الاجمالى'] + merge_df['اجمالى اليوم']
        final_df = merge_df[['المنتجات','العدد الكلى','اجمالى اليوم']]
        
        final_df.to_csv(ta2rer_today,index=False)
        

        
        save_docx = []
        
        with open(file_name,'r',encoding='utf-8') as file6:
            write2 = csv.DictReader(file6)
            for row in write2:
                save_docx.append(row)
                
        out_docx = []
        for row in save_docx:
            total_mohamed = row['الاجمالى']
        out_docx1 = []
        if not os.path.exists(file_name1):
            montag = []
            with open('products.csv', 'r', newline='', encoding='utf-8') as file:
                reader = csv.DictReader(file)
                for row in reader:
                    montag.append(row['prodects'])
            with open(file_name1,'w',encoding='utf-8') as file9:
                write5 = csv.DictWriter(file9,fieldnames=['المنتجات','العدد الكلى','الاجمالى'])
                write5.writerow({'المنتجات':'المنتجات','العدد الكلى':'العدد الكلى','الاجمالى':'اجمالى اليوم'})
                for row in montag:
                    write5.writerow({'المنتجات':row})
        with open(file_name1,'r',encoding='utf-8') as file7:
            write2 = csv.DictReader(file7)
            for row in write2:
                out_docx1.append(row)

        number = []

            
        montag = []
        with open('products.csv', 'r', newline='', encoding='utf-8') as file:
            reader = csv.DictReader(file)
            for row in reader:
                montag.append(row['prodects'])
        
        with open(file_name1,'w',encoding='utf-8') as file8:
            writer1 = csv.DictWriter(file8,fieldnames=['المنتجات','العدد','الاجمالى'])
            writer1.writerow({'المنتجات':'المنتجات','العدد':'العدد الكلى','الاجمالى':'اجمالى اليوم كله'})
            for row in montag:
                writer1.writerow({'المنتجات':row})
                
        
            
                    
                    
    def every_day():
        today_date = datetime.datetime.now().strftime("%Y-%m-%d")
        today_date1 = datetime.datetime.now().strftime("%Y-%m")
        
        folder_path = os.path.join(os.getcwd(),'daily',today_date1, today_date)
        if not os.path.exists(folder_path):
            messagebox.showinfo('error','لا يوجد مبيعات لهذا اليوم')
            
            return
        
        sales_data = defaultdict(float)
        files = [file for file in os.listdir(folder_path) if file.endswith(".txt")]
        for file_name in files:
            file_path = os.path.join(folder_path, file_name)
            with open(file_path, "r",encoding='utf-8') as file:
                for line in file:
                    # product, price = line.strip().split(":")
                    # sales_data[product] += float(price)
                    line1 = line.split()
                    # print(line1[2])
                    
        for product, total_sales in sales_data.items():
            Label(root, text=f"{product}:{ total_sales}").pack()
            
        root_ta2rer_today = Tk()
        root_ta2rer_today.geometry('480x580')
        root_ta2rer_today.title('التقرير اليومى')
        light = []
        file_name1 = os.path.join(folder_path, "ta2rer.csv")
        ta2rer_today = os.path.join(folder_path, "ta2rer_today.csv")

        with open(ta2rer_today,'r',newline='',encoding='utf-8') as file:
            row_number = 1
            reader = csv.DictReader(file)
            for row in reader:
                light.append(row)
                product_name = row['المنتجات']  # استخدم اسم العمود الصحيح هنا
                price = row['العدد الكلى']  # استخدم اسم العمود الصحيح هنا
                total_today = row['اجمالى اليوم']  # استخدم اسم العمود الصحيح هنا
        
                button = Label(root_ta2rer_today, text=product_name)
                button.grid(row=row_number, column=0, padx=3, pady=3)
                
                label1 = Label(root_ta2rer_today,text=price)
                # entry.insert(0, price)
                label1.grid(row=row_number, column=1, padx=3, pady=3)
                # entries[product_name] = entry
                row_number += 1
            label2 = Label(root_ta2rer_today,text='الاجمالى')
            label2.grid()
            label3 = Label(root_ta2rer_today, text=total_today)
            label3.grid(row=row_number, column=1)

        
        root_ta2rer_today.mainloop()
            
    def total_money2(event):
        # total_fatora1.insert(0,END)
        total_get = total_fatora.get("1.0", "end-1c")
        total_get1 = float(total_get)
        total_get3 = float(total_fatora1.get())
        total_get2 = total_get3 - total_get1
        total_fatora2.delete('1.0',END)
        total_fatora2.insert(END,f'{total_get2}')
        

    def total_money(text):
        if text.isdigit() or text == "":
            
            # total_get3 = float(total_fatora1.get())
            total_fatora1.bind('<KeyRelease>',total_money2)
            return True
        else:
            messagebox.showerror("Error", "Please enter only numbers.")
            return False





    root = Tk()
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    root.geometry(f"{screen_width}x{screen_height}")
    root.resizable(False,False)
    root.title('كاشير محلات حواوشى الليثى')
    title = Label(root,text='حواوشى الليثى',fg='gold',bg='black',font=('tajwal',16,'bold'))
    title.pack(fill='x')

    left1 = Frame(root, width=300,height=1000,borderwidth=5,relief='ridge')
    left1.place(x=0,y=35)
    title1 = Label(left1,text='الماكولات',height=1,width=34,fg='gold',bg='black',font=('tajwal',16,'bold'))
    title1.grid(row=0,column=0,columnspan=3)
    frame = Frame(root)
    frame.pack()

    # قراءة الأزرار والأسعار من ملف النص

    # إنشاء الأزرار ومربعات الإدخال
    entries = {}


    def load_data():
        with open('products.csv', 'r', newline='', encoding='utf-8') as file:
            row_number = 1
            reader = csv.DictReader(file)
            for row in reader:
                product_name = row['prodects']  # استخدم اسم العمود الصحيح هنا
                price = row['price']  # استخدم اسم العمود الصحيح هنا
                
                button = Label(left1, text=product_name)
                button.grid(row=row_number, column=0, padx=10, pady=5)
                
                entry = Entry(left1)
                # entry.insert(0, price)
                entry.grid(row=row_number, column=1, padx=10, pady=5)
                entries[product_name] = entry
                button = Button(left1, text="اضافه", command=lambda n=product_name: multiply_price(n))
                button.grid(row=row_number, column=2)
                
                row_number += 1
                
    load_data()
    right = Frame(root,width=80,height=80,borderwidth=5,relief='ridge')
    right.place(x=screen_width-360,y=screen_height -160)

    button_calc = Button(right,text='تفريغ الخانات',height=3,width=15,command=fatora_empty)
    button_calc.grid(row=0,column=0)
    button_empty = Button(right,text='التقرير اليومى',height=3,width=15,command=every_day)
    button_empty.grid(row=0,column=1)
    button_save = Button(right,text='حفظ وطباعه الفاتوره',height=3,width=15,command=save_fatora)
    button_save.grid(row=1,column=0)
    button_close = Button(right,text='اغلاق البرنامج',height=3,width=15,command=close)
    button_close.grid(row=1,column=1)

    qur_width = screen_width // 4
    qur_height = screen_height // 5

    up_right = Frame(root,width=qur_width,height=qur_height,borderwidth=5,relief='ridge')
    up_right.place(x=screen_width-445,y=screen_height-735)

    fatora_text = Text(up_right,width=50,height=1,bg='green')
    fatora_text.tag_configure("center", justify='center')
    fatora_text.tag_add("center", "1.0", "end")
    fatora_text.insert(END,'     السعر\t        الكميه \t            \t الصنف\n')
    fatora_text.pack(fill=BOTH,expand=1)
    fatora_listbox = Listbox(up_right,width=70,height=20)
    fatora_listbox.pack(fill=BOTH,expand=1)
    up_right_1 = Frame(up_right,height=10,width=60)
    up_right_1.pack()
    total_fatora = Text(up_right_1,width=30,height=1)
    total_fatora.insert(END,0)
    total_fatora.tag_configure("center", justify='center')
    total_fatora.tag_add("center", "1.0", "end")
    total_fatora.grid(row=0,column=0)
    all_fatora = Label(up_right_1,text='الاجمالى')
    all_fatora.grid(row=0,column=1)
    up_right_2 = Frame(up_right,height=10,width=60)
    up_right_2.pack()
    validate_cmd = up_right_2.register(total_money)
    total_fatora1 = Entry(up_right_2,width=40,validate='key',validatecommand=(validate_cmd, "%P"),justify='center')
    # total_fatora1.insert(0,0)
    total_fatora1.grid(row=0,column=0)
    all_fatora1 = Label(up_right_2,text='المدفوع')
    all_fatora1.grid(row=0,column=1)
    up_right_3 = Frame(up_right,height=10,width=60)
    up_right_3.pack()
    total_fatora2 = Text(up_right_3,width=30,height=1)
    total_fatora2.insert(END,0)
    total_fatora2.tag_configure("center", justify='center')
    total_fatora2.tag_add("center", "1.0", "end")
    total_fatora2.grid(row=0,column=0)
    all_fatora2 = Label(up_right_3,text='  الباقى')
    all_fatora2.grid(row=0,column=1)
    Listbox_delet = Button(up_right,text='حذف',height=3,width=60,command=delete_food1)
    Listbox_delet.pack()


    root.mainloop()
seller1()